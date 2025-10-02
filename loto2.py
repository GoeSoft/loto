import random
import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Диапазоны для 5 столбцов
COLUMN_RANGES = [
    list(range(1, 10)),    # 1-9
    list(range(10, 20)),   # 10-19
    list(range(20, 30)),   # 20-29
    list(range(30, 40)),   # 30-39
    list(range(40, 50))    # 40-49
]

# Глобальная переменная для хранения последних билетов
last_tickets = []

def generate_ticket():
    empty_cols = [random.randint(0, 4) for _ in range(3)]
    grid = [[None for _ in range(5)] for _ in range(3)]

    for col in range(5):
        non_empty_rows = [row for row in range(3) if empty_cols[row] != col]
        count_needed = len(non_empty_rows)
        if count_needed == 0:
            continue
        numbers = random.sample(COLUMN_RANGES[col], count_needed)
        random.shuffle(numbers)
        for i, row in enumerate(non_empty_rows):
            grid[row][col] = numbers[i]

    return tuple(tuple(row) for row in grid)

def generate_all_tickets(n=40):
    tickets = set()
    attempts = 0
    max_attempts = n * 100

    while len(tickets) < n and attempts < max_attempts:
        ticket = generate_ticket()
        tickets.add(ticket)
        attempts += 1

    return list(tickets)

def format_ticket_for_text(ticket, idx):
    lines = []
    lines.append(f"Билет #{idx + 1}")
    lines.append("-" * 25)
    for row in ticket:
        cells = []
        for cell in row:
            if cell is None:
                cells.append(" --")
            else:
                cells.append(f"{cell:3d}")
        lines.append(" ".join(cells))
    lines.append("-" * 25)
    return "\n".join(lines)

def format_ticket_for_word(ticket, idx):
    """Возвращает строки для Word (без дефисов, но с отступами)"""
    lines = []
    lines.append(f"Билет #{idx + 1}")
    for row in ticket:
        cells = []
        for cell in row:
            if cell is None:
                cells.append("--")
            else:
                cells.append(f"{cell:2d}")  # 2 символа
        # Соединяем через табуляцию или пробелы
        lines.append("   ".join(cells))
    return lines

def on_generate():
    global last_tickets
    try:
        last_tickets = generate_all_tickets(40)
        output_text.delete(1.0, tk.END)
        for i, t in enumerate(last_tickets):
            output_text.insert(tk.END, format_ticket_for_text(t, i) + "\n\n")
        output_text.see(tk.END)
    except Exception as e:
        messagebox.showerror("Ошибка", str(e))

def on_save_to_word():
    global last_tickets
    if not last_tickets:
        messagebox.showwarning("Внимание", "Сначала сгенерируйте билеты!")
        return

    # Выбор пути сохранения
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word документ", "*.docx")],
        title="Сохранить билеты в Word"
    )
    if not file_path:
        return  # отмена

    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(11)

        # Заголовок
        title = doc.add_heading('Сгенерированные лото-билеты', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, ticket in enumerate(last_tickets):
            # Добавляем "Билет #1"
            doc.add_heading(f"Билет #{i + 1}", level=2)
            for line in format_ticket_for_word(ticket, i):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Пустой абзац между билетами
            doc.add_paragraph()

        doc.save(file_path)
        messagebox.showinfo("Успех", f"Билеты сохранены в:\n{file_path}")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{str(e)}")

# Создаём окно
root = tk.Tk()
root.title("Генератор лото-билетов")
root.geometry("800x600")

# Кнопки
btn_frame = tk.Frame(root)
btn_frame.pack(pady=10)

btn_generate = tk.Button(btn_frame, text="Сгенерировать 40 билетов", command=on_generate, font=("Arial", 12), bg="#4CAF50", fg="white")
btn_generate.pack(side=tk.LEFT, padx=5)

btn_save = tk.Button(btn_frame, text="Сохранить в Word", command=on_save_to_word, font=("Arial", 12), bg="#2196F3", fg="white")
btn_save.pack(side=tk.LEFT, padx=5)

# Поле вывода
output_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, font=("Courier", 10), bg="white", fg="black")
output_text.pack(expand=True, fill=tk.BOTH, padx=10, pady=10)

# Запуск
root.mainloop()